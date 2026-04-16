"""Report Generator - Terminal, Excel, Word, and DQR PPT output.

Generates QC reports in multiple formats from QC analysis results.
"""

from __future__ import annotations

import datetime
import sys
from pathlib import Path
from dataclasses import dataclass

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_WORD = True
except ImportError:
    HAS_WORD = False

try:
    import sys as _sys
    from pathlib import Path as _Path
    _shared = _Path(__file__).resolve().parents[3] / "_shared"
    if _shared.exists() and str(_shared) not in _sys.path:
        _sys.path.insert(0, str(_shared))
    from geoview_common.reporting.design_system import WordBuilder, WORD_STYLES
    HAS_WORDBUILDER = True
except ImportError:
    HAS_WORDBUILDER = False


# ── Color definitions ───────────────────────────────────────────

_GREEN = "00B050"
_RED = "FF0000"
_YELLOW = "FFC000"
_GRAY = "808080"


def _status_color(status: str) -> str:
    return {"PASS": _GREEN, "FAIL": _RED, "WARNING": _YELLOW}.get(status, _GRAY)


def _safe_print(*args, **kwargs) -> None:
    """Print using the active console encoding without crashing on cp949."""
    stream = sys.stdout or sys.__stdout__
    if stream is None:
        return
    sep = kwargs.get("sep", " ")
    end = kwargs.get("end", "\n")
    text = sep.join(str(arg) for arg in args) + end
    encoding = getattr(stream, "encoding", None) or "utf-8"
    safe = text.encode(encoding, errors="replace").decode(encoding, errors="replace")
    stream.write(safe)


def _sanitize_report_text(text: str | None, max_lines: int = 4, max_chars: int = 320) -> str:
    value = str(text or "").replace("\r", "")
    for old, new in (("°", " deg"), ("±", "+/-"), ("→", "->"), ("—", "-"), ("–", "-")):
        value = value.replace(old, new)
    lines = [" ".join(line.strip().split()) for line in value.split("\n") if line.strip()]
    if not lines:
        return ""
    extra = max(0, len(lines) - max_lines)
    lines = lines[:max_lines]
    compact = " / ".join(lines)
    if extra:
        compact += f" / ... {extra} more lines"
    if len(compact) > max_chars:
        compact = compact[: max_chars - 3].rstrip() + "..."
    return compact


def _extract_items(result) -> list[dict]:
    items = []
    if hasattr(result, "items"):
        raw_items = result.items if isinstance(result.items, list) else []
        for item in raw_items:
            if isinstance(item, dict):
                items.append(item)
            elif hasattr(item, "name"):
                items.append({
                    "name": getattr(item, "name", ""),
                    "status": getattr(item, "status", "N/A"),
                    "detail": getattr(item, "detail", ""),
                })
    elif hasattr(result, "checks"):
        for item in getattr(result, "checks", []):
            detail_parts = []
            if getattr(item, "pds_value", ""):
                detail_parts.append(f"PDS={getattr(item, 'pds_value', '')}")
            if getattr(item, "reference_value", ""):
                detail_parts.append(f"Ref={getattr(item, 'reference_value', '')}")
            if getattr(item, "difference", ""):
                detail_parts.append(f"Diff={getattr(item, 'difference', '')}")
            if getattr(item, "suggestion", ""):
                detail_parts.append(f"Action={getattr(item, 'suggestion', '')}")
            items.append({
                "name": f"{getattr(item, 'category', '')} / {getattr(item, 'name', '')}".strip(" /"),
                "status": getattr(item, "status", "N/A"),
                "detail": " | ".join(detail_parts),
            })
    elif hasattr(result, "roll_bias_deg") or hasattr(result, "pitch_bias_deg"):
        items.append({
            "name": "Roll Bias",
            "status": getattr(result, "roll_verdict", "N/A"),
            "detail": (
                f"{getattr(result, 'roll_bias_deg', 0.0):+.4f} deg +/- "
                f"{getattr(result, 'roll_bias_std', 0.0):.4f} deg "
                f"({getattr(result, 'roll_num_pings', 0)} pings)"
            ),
        })
        items.append({
            "name": "Pitch Bias",
            "status": getattr(result, "pitch_verdict", "N/A"),
            "detail": (
                f"{getattr(result, 'pitch_bias_deg', 0.0):+.4f} deg +/- "
                f"{getattr(result, 'pitch_bias_std', 0.0):.4f} deg "
                f"({getattr(result, 'pitch_num_pairs', 0)} pairs)"
            ),
        })
        hvf_vs_data = getattr(result, "hvf_vs_data", "")
        if hvf_vs_data:
            items.append({
                "name": "HVF vs Data",
                "status": _extract_verdict(result),
                "detail": hvf_vs_data,
            })
    elif hasattr(result, "gap_verdict") or hasattr(result, "roll"):
        for axis_name in ("roll", "pitch", "heave", "heading"):
            axis = getattr(result, axis_name, None)
            if not axis:
                continue
            items.append({
                "name": getattr(axis, "name", axis_name.title()),
                "status": getattr(axis, "verdict", "N/A"),
                "detail": (
                    f"mean={getattr(axis, 'mean', 0.0):+.4f}{getattr(axis, 'unit', '')}, "
                    f"std={getattr(axis, 'std', 0.0):.4f}{getattr(axis, 'unit', '')}, "
                    f"spikes={getattr(axis, 'num_spikes', 0)} "
                    f"({getattr(axis, 'spike_rate_pct', 0.0):.2f}%)"
                ),
            })
        items.append({
            "name": "Gap Analysis",
            "status": getattr(result, "gap_verdict", "N/A"),
            "detail": (
                f"{getattr(result, 'num_gaps', 0)} gaps, "
                f"max {getattr(result, 'max_gap_sec', 0.0):.3f}s, "
                f"sample rate {getattr(result, 'sample_rate_hz', 0.0):.1f}Hz"
            ),
        })
    for item in items:
        item["name"] = _sanitize_report_text(item.get("name", ""), max_lines=1, max_chars=96)
        item["detail"] = _sanitize_report_text(item.get("detail", ""))
    return items


def _extract_verdict(result) -> str:
    if hasattr(result, "overall_verdict"):
        verdict = getattr(result, "overall_verdict")
        if callable(verdict):
            verdict = verdict()
        return str(verdict or "N/A")
    if hasattr(result, "overall"):
        return str(getattr(result, "overall") or "N/A")
    if hasattr(result, "verdict"):
        return str(getattr(result, "verdict") or "N/A")

    statuses = []
    for item in _extract_items(result):
        status = str(item.get("status", "N/A")).upper()
        if status not in ("", "N/A", "INFO"):
            statuses.append(status)
    if "FAIL" in statuses:
        return "FAIL"
    if "WARNING" in statuses:
        return "WARNING"
    return "PASS" if statuses else "N/A"


def _aggregate_overall_verdict(verdicts: list[str]) -> str:
    overall = "N/A"
    seen_real_verdict = False
    for verdict in verdicts:
        verdict_text = str(verdict or "").upper()
        if verdict_text not in ("PASS", "WARNING", "FAIL"):
            continue
        seen_real_verdict = True
        if verdict_text == "FAIL":
            return "FAIL"
        if verdict_text == "WARNING":
            overall = "WARNING"
        elif overall == "N/A":
            overall = "PASS"
    return overall if seen_real_verdict else "N/A"


# ── Excel Report ────────────────────────────────────────────────


def generate_excel_report(
    qc_results: dict,
    output_path: str | Path,
    project_name: str = "",
    vessel_name: str = "",
) -> None:
    """Generate Excel QC report with professional styling."""
    if not HAS_EXCEL:
        print("  [SKIP] openpyxl not installed")
        return

    wb = openpyxl.Workbook()

    # ── Style constants ──
    NAVY = "1E3A5F"
    navy_fill = PatternFill("solid", fgColor=NAVY)
    hdr_font = Font(bold=True, size=11, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center")
    title_font = Font(bold=True, size=16, color="FFFFFF")
    title_align = Alignment(horizontal="center", vertical="center")
    info_label_font = Font(bold=True, size=10, color=NAVY)
    info_value_font = Font(size=10)
    data_font = Font(size=10)
    alt_fill = PatternFill("solid", fgColor="F2F6FA")
    pass_fill = PatternFill("solid", fgColor="C6EFCE")
    fail_fill = PatternFill("solid", fgColor="FFC7CE")
    warn_fill = PatternFill("solid", fgColor="FFEB9C")
    pass_font = Font(bold=True, size=10, color="006100")
    fail_font = Font(bold=True, size=10, color="9C0006")
    warn_font = Font(bold=True, size=10, color="9C6500")
    thin_border = Border(
        left=Side(style="thin", color="D0D0D0"),
        right=Side(style="thin", color="D0D0D0"),
        top=Side(style="thin", color="D0D0D0"),
        bottom=Side(style="thin", color="D0D0D0"),
    )

    def _status_style(val: str):
        """Return (font, fill) for a status value."""
        if val == "PASS":
            return pass_font, pass_fill
        elif val == "FAIL":
            return fail_font, fail_fill
        elif val == "WARNING":
            return warn_font, warn_fill
        return data_font, None

    def _add_sheet(name: str, headers: list[str], rows: list[list],
                   start_row: int = 1) -> None:
        ws = wb.create_sheet(name)

        # Header row
        for c, h in enumerate(headers, 1):
            cell = ws.cell(row=start_row, column=c, value=h)
            cell.font = hdr_font
            cell.fill = navy_fill
            cell.alignment = hdr_align
            cell.border = thin_border

        # Data rows with alternating fills & status cell coloring
        for r, row_data in enumerate(rows, start_row + 1):
            is_alt = (r - start_row) % 2 == 0
            for c, val in enumerate(row_data, 1):
                cell = ws.cell(row=r, column=c, value=str(val))
                cell.border = thin_border
                cell.font = data_font
                # Alternating row background
                if is_alt:
                    cell.fill = alt_fill

                # Status cells: colored fill + bold font
                if isinstance(val, str) and val in ("PASS", "FAIL", "WARNING"):
                    sfont, sfill = _status_style(val)
                    cell.font = sfont
                    if sfill:
                        cell.fill = sfill

        # Auto-width
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(50, max(12, max_len + 2))

        return ws

    # ── Summary sheet ──
    ws_sum = wb.create_sheet("Summary", 0)

    # Navy title banner (merged A1:F1)
    ws_sum.merge_cells("A1:F1")
    title_cell = ws_sum["A1"]
    title_cell.value = "MBES QC Report"
    title_cell.font = title_font
    title_cell.fill = navy_fill
    title_cell.alignment = title_align
    ws_sum.row_dimensions[1].height = 36

    # Project info block (rows 3-6)
    info_items = [
        ("Project", project_name or "N/A"),
        ("Vessel", vessel_name or "N/A"),
        ("Report Date", datetime.datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]
    for idx, (label, value) in enumerate(info_items):
        r = idx + 3
        lbl_cell = ws_sum.cell(row=r, column=1, value=label)
        lbl_cell.font = info_label_font
        val_cell = ws_sum.cell(row=r, column=2, value=value)
        val_cell.font = info_value_font

    # QC Results header
    qc_start_row = len(info_items) + 4
    summary_rows = []
    verdicts = []
    for category, result in qc_results.items():
        verdict = _extract_verdict(result)

        summary_rows.append([category, verdict])
        verdicts.append(verdict)

    overall = _aggregate_overall_verdict(verdicts)

    summary_rows.insert(0, ["OVERALL", overall])

    # Write summary table headers
    for c, h in enumerate(["Category", "Verdict"], 1):
        cell = ws_sum.cell(row=qc_start_row, column=c, value=h)
        cell.font = hdr_font
        cell.fill = navy_fill
        cell.alignment = hdr_align
        cell.border = thin_border

    # Write summary data
    for r_idx, row_data in enumerate(summary_rows, qc_start_row + 1):
        is_alt = (r_idx - qc_start_row) % 2 == 0
        for c, val in enumerate(row_data, 1):
            cell = ws_sum.cell(row=r_idx, column=c, value=str(val))
            cell.border = thin_border
            cell.font = data_font
            if is_alt:
                cell.fill = alt_fill
            if isinstance(val, str) and val in ("PASS", "FAIL", "WARNING"):
                sfont, sfill = _status_style(val)
                cell.font = sfont
                if sfill:
                    cell.fill = sfill

    ws_sum.column_dimensions["A"].width = 25
    ws_sum.column_dimensions["B"].width = 20

    # ── Detail sheets per category ──
    for category, result in qc_results.items():
        items = [
            [item.get("name", ""), item.get("status", ""), item.get("detail", "")]
            for item in _extract_items(result)
        ]

        if items:
            safe_name = category[:31].replace("/", "-").replace("\\", "-")
            _add_sheet(safe_name, ["Check", "Status", "Detail"], items)

    # Remove default sheet
    if "Sheet" in wb.sheetnames:
        del wb["Sheet"]

    wb.save(str(output_path))


# ── Word Report ─────────────────────────────────────────────────


def generate_word_report(
    qc_results: dict,
    output_path: str | Path,
    project_name: str = "",
    vessel_name: str = "",
) -> None:
    """Generate Word QC report using GeoView design system."""
    if not HAS_WORDBUILDER:
        if not HAS_WORD:
            print("  [SKIP] python-docx not installed")
            return
        # Fallback: minimal raw-docx report
        _generate_word_report_fallback(qc_results, output_path, project_name, vessel_name)
        return

    wb = WordBuilder(WORD_STYLES["geoview_report"])

    # ── Cover Page ──
    wb.cover(
        title="MBES QC Report",
        subtitle="Multi-Beam Echo Sounder Quality Control",
        meta=f"Project: {project_name or 'N/A'}  |  "
             f"Vessel: {vessel_name or 'N/A'}  |  "
             f"Date: {datetime.datetime.now():%Y-%m-%d}",
    )
    wb.setup_page_footer("MBES QC Report")
    wb.page_break()

    # ── 1. Project Information ──
    wb.reset_numbering()
    wb.heading("Project Information", level=1)
    wb.kv_table([
        ("Project", project_name or "N/A"),
        ("Vessel", vessel_name or "N/A"),
        ("Report Date", datetime.datetime.now().strftime("%Y-%m-%d %H:%M")),
    ])

    # ── 2. QC Summary ──
    wb.heading("QC Summary", level=1)

    summary_rows = []
    verdicts = []
    for category, result in qc_results.items():
        verdict = _extract_verdict(result)

        summary_rows.append([category, verdict])
        verdicts.append(verdict)

    overall = _aggregate_overall_verdict(verdicts)

    summary_rows.insert(0, ["OVERALL", overall])
    wb.table(
        headers=["Category", "Verdict"],
        rows=summary_rows,
    )

    # ── 3. Detailed Results ──
    wb.heading("Detailed Results", level=1)

    for category, result in qc_results.items():
        wb.heading(category, level=2)

        items = []
        items = _extract_items(result)

        if items:
            detail_rows = []
            for item in items:
                detail_rows.append([
                    item.get("name", ""),
                    item.get("status", ""),
                    item.get("detail", ""),
                ])
            wb.table(
                headers=["Check", "Status", "Detail"],
                rows=detail_rows,
            )
        else:
            wb.body_text("No detailed check items available for this category.")

    wb.doc.save(str(output_path))


def _generate_word_report_fallback(
    qc_results: dict, output_path: str | Path,
    project_name: str, vessel_name: str,
) -> None:
    """Minimal fallback Word report when WordBuilder is unavailable."""
    doc = Document()
    title = doc.add_heading("MBES QC Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run("Project: ").bold = True
    p.add_run(project_name or "N/A")
    p = doc.add_paragraph()
    p.add_run("Vessel: ").bold = True
    p.add_run(vessel_name or "N/A")
    p = doc.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    for category, result in qc_results.items():
        doc.add_heading(category, level=2)
    doc.save(str(output_path))


# ── Terminal Report ─────────────────────────────────────────────


def print_terminal_report(qc_results: dict) -> None:
    """Print colored terminal QC report."""

    def _vc(v: str) -> str:
        c = {"PASS": "\033[92m", "WARNING": "\033[93m", "FAIL": "\033[91m"}.get(v, "")
        return f"{c}{v}\033[0m"

    _safe_print("\n" + "=" * 70)
    _safe_print("  MBES QC REPORT")
    _safe_print("=" * 70)

    verdicts = []
    for category, result in qc_results.items():
        verdict = _extract_verdict(result)
        verdicts.append(verdict)

        _safe_print(f"\n  [{_vc(verdict)}] {category}")

        for item in _extract_items(result):
            s = item.get("status", "N/A")
            _safe_print(f"      {_vc(s):>20s}  {item.get('name', '')}: {item.get('detail', '')}")

    overall = _aggregate_overall_verdict(verdicts)

    _safe_print("\n" + "=" * 70)
    _safe_print(f"  OVERALL: {_vc(overall)}")
    _safe_print("=" * 70 + "\n")
